"""Format-aware policy document extraction.

Supports PDF (font/block metadata when available), DOCX (heading styles),
Markdown headings, and plain text. Nothing here is company-specific —
rule IDs, categories, and remedy text all come from the document itself.
"""

from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass, field
from pathlib import Path

from src import llm_client, prompts
from src.schema import PolicyRule

RULE_ID_RE = re.compile(
    r"^(?P<id>(?:[A-Z]\d+(?:\.\d+)*)|(?:\d+(?:\.\d+)+)|(?:Section\s+\d+))\b[^\n]*",
    re.IGNORECASE,
)
HEADING_MD_RE = re.compile(r"^(#{1,3})\s+(.+)$", re.MULTILINE)
PARAGRAPH_SPLIT = re.compile(r"\n\s*\n")


@dataclass
class RawSection:
    """One extracted section before rule normalization."""

    heading: str
    body: str
    source: str = "native"  # native | llm
    meta: dict = field(default_factory=dict)

    @property
    def text(self) -> str:
        return f"{self.heading}\n{self.body}".strip() if self.heading else self.body.strip()

    @property
    def content_hash(self) -> str:
        return hashlib.sha256(self.text.encode("utf-8")).hexdigest()


def section_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def extract_sections(path: str | Path, *, use_llm_fallback: bool = True) -> list[RawSection]:
    """Extract sections from a policy document. Format inferred from suffix."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        sections = _extract_pdf(path)
    elif suffix == ".docx":
        sections = _extract_docx(path)
    elif suffix in (".md", ".markdown"):
        sections = _extract_markdown(path.read_text(encoding="utf-8", errors="replace"))
    elif suffix in (".txt", ".text"):
        sections = _extract_plain(path.read_text(encoding="utf-8", errors="replace"))
    else:
        # Unknown extension: try PDF then plain text.
        try:
            sections = _extract_pdf(path)
        except Exception:
            sections = _extract_plain(path.read_text(encoding="utf-8", errors="replace"))

    if len(sections) < 2 and use_llm_fallback:
        raw = path.read_bytes()
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = _pdf_plain_text(path) if suffix == ".pdf" else ""
        if text.strip():
            llm_sections = _llm_segment(text)
            if llm_sections:
                return llm_sections
    return sections


def normalize_rule(section: RawSection, *, use_llm: bool = False) -> PolicyRule:
    """Turn a raw section into a structured PolicyRule."""
    heading, body = section.heading.strip(), section.body.strip()
    full = section.text
    rule_id = _guess_rule_id(heading) or _guess_rule_id(full) or _slug_id(heading or full[:40])

    condition, outcome = _split_condition_outcome(body or full)
    category = _guess_category(heading, body)
    region = ""
    effective_date = ""

    if use_llm and (not condition or not outcome or category == "global"):
        parsed = _llm_normalize(full)
        if parsed:
            rule_id = str(parsed.get("id") or rule_id).strip() or rule_id
            condition = str(parsed.get("condition") or condition).strip()
            outcome = str(parsed.get("outcome") or outcome).strip()
            category = str(parsed.get("category") or category).strip() or "global"
            region = str(parsed.get("region") or "").strip()
            effective_date = str(parsed.get("effective_date") or "").strip()

    # Cross-cutting escalation / override language stays in every scoped search.
    if _looks_global(heading, body):
        category = "global"

    depends_on, dependency_status = _infer_depends_on(condition, outcome, full)
    return PolicyRule(
        id=rule_id,
        condition=condition,
        outcome=outcome,
        category=category.lower().replace(" ", "_") or "global",
        region=region,
        effective_date=effective_date,
        text=full,
        section_hash=section.content_hash,
        depends_on=depends_on,
        dependency_status=dependency_status,
    )


def normalize_sections(
    sections: list[RawSection],
    *,
    changed_hashes: set[str] | None = None,
    use_llm_for_changed: bool = True,
) -> list[PolicyRule]:
    """Normalize all sections; optionally LLM-refine only changed ones."""
    rules = []
    for sec in sections:
        need_llm = bool(
            use_llm_for_changed
            and changed_hashes is not None
            and sec.content_hash in changed_hashes
        )
        rules.append(normalize_rule(sec, use_llm=need_llm))
    return _dedupe_ids(rules)


# --------------------------------------------------------------------------- extractors


def _extract_pdf(path: Path) -> list[RawSection]:
    """Prefer PyMuPDF block+font metadata; fall back to pypdf plain text."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(path)
        blocks: list[tuple[float, str, bool]] = []
        for page in doc:
            for b in page.get_text("dict").get("blocks", []):
                if b.get("type") != 0:
                    continue
                lines = []
                sizes = []
                flags = []
                for line in b.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    text = "".join(s.get("text", "") for s in spans).strip()
                    if text:
                        lines.append(text)
                        sizes.extend(s.get("size", 10) for s in spans)
                        flags.extend(s.get("flags", 0) for s in spans)
                if not lines:
                    continue
                text = " ".join(lines)
                avg_size = sum(sizes) / len(sizes) if sizes else 10
                bold = any(f & 2**4 for f in flags)  # bit 4 = bold in MuPDF
                blocks.append((avg_size, text, bold))
        doc.close()
        if not blocks:
            return _extract_plain(_pdf_plain_text(path))
        median = sorted(s for s, _, _ in blocks)[len(blocks) // 2]
        sections: list[RawSection] = []
        cur_h, cur_b = "", []
        for size, text, bold in blocks:
            is_heading = (bold and size >= median) or RULE_ID_RE.match(text) or size >= median + 1.5
            if is_heading and len(text) < 200:
                if cur_h or cur_b:
                    sections.append(RawSection(cur_h, "\n".join(cur_b), source="native"))
                cur_h, cur_b = text, []
            else:
                if not cur_h and not cur_b and RULE_ID_RE.match(text):
                    cur_h = text
                else:
                    cur_b.append(text)
        if cur_h or cur_b:
            sections.append(RawSection(cur_h, "\n".join(cur_b), source="native"))
        # Drop title-only fluff without a body when we have enough real rules.
        sections = [s for s in sections if len(s.text) > 40]
        # Drop title-only / preamble blocks that aren't actual rules.
        sections = [s for s in sections if _looks_like_rule(s)]
        if len(sections) >= 2:
            return sections
    except Exception:
        pass
    return _extract_plain(_pdf_plain_text(path))


def _pdf_plain_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(path: Path) -> list[RawSection]:
    from docx import Document

    doc = Document(str(path))
    sections: list[RawSection] = []
    cur_h, cur_b = "", []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = (p.style.name or "") if p.style else ""
        is_heading = style.startswith("Heading") or bool(RULE_ID_RE.match(text))
        if is_heading:
            if cur_h or cur_b:
                sections.append(RawSection(cur_h, "\n".join(cur_b), source="native"))
            cur_h, cur_b = text, []
        else:
            cur_b.append(text)
    if cur_h or cur_b:
        sections.append(RawSection(cur_h, "\n".join(cur_b), source="native"))
    sections = [s for s in sections if len(s.text) > 40]
    return sections if sections else _extract_plain("\n\n".join(p.text for p in doc.paragraphs if p.text))


def _extract_markdown(text: str) -> list[RawSection]:
    matches = list(HEADING_MD_RE.finditer(text))
    if not matches:
        return _extract_plain(text)
    sections: list[RawSection] = []
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        heading = m.group(2).strip()
        if len(heading) + len(body) > 40:
            sections.append(RawSection(heading, body, source="native"))
    return sections or _extract_plain(text)


def _extract_plain(text: str) -> list[RawSection]:
    pattern = re.compile(r"\n(?=(?:[A-Z]?\d+(?:\.\d+)?\s)|(?:Section\s+\d+))", re.IGNORECASE)
    parts = [p.strip() for p in pattern.split(text) if len(p.strip()) > 40]
    if len(parts) < 3:
        parts = [p.strip() for p in PARAGRAPH_SPLIT.split(text) if len(p.strip()) > 40]
    sections = []
    for p in parts:
        lines = p.splitlines()
        heading = lines[0].strip() if lines and RULE_ID_RE.match(lines[0].strip()) else ""
        body = "\n".join(lines[1:]).strip() if heading else p
        if not heading:
            heading = lines[0].strip()[:120] if lines else "section"
            body = "\n".join(lines[1:]).strip() or p
        sections.append(RawSection(heading, body, source="native"))
    return sections


def _llm_segment(text: str) -> list[RawSection]:
    """Last-resort: ask the LLM to split an unstructured document into sections."""
    try:
        raw = llm_client.complete(
            prompts.POLICY_SEGMENT_SYSTEM,
            prompts.POLICY_SEGMENT_USER.format(document=text[:12000]),
            max_tokens=2000,
            purpose="generate",
        )
        data = llm_client.extract_json(raw)
        items = data.get("sections") if isinstance(data, dict) else None
        if not isinstance(items, list):
            return []
        out = []
        for item in items:
            if not isinstance(item, dict):
                continue
            h = str(item.get("heading") or "").strip()
            b = str(item.get("body") or "").strip()
            if len(h) + len(b) > 40:
                out.append(RawSection(h, b, source="llm"))
        return out
    except Exception:
        return []


def _llm_normalize(text: str) -> dict | None:
    try:
        raw = llm_client.complete(
            prompts.POLICY_NORMALIZE_SYSTEM,
            prompts.POLICY_NORMALIZE_USER.format(section=text[:4000]),
            max_tokens=600,
            purpose="generate",
        )
        data = llm_client.extract_json(raw)
        return data if isinstance(data, dict) else None
    except Exception:
        return None


# --------------------------------------------------------------------------- helpers


def _looks_like_rule(section: RawSection) -> bool:
    """Skip document titles / preambles that aren't actionable rules."""
    text = section.text
    if RULE_ID_RE.match(section.heading.strip()) or RULE_ID_RE.match(text.strip()):
        return True
    if re.search(r"\bMUST\b|\bMUST NOT\b|\bshall\b", text, re.IGNORECASE):
        return True
    # Keep substantial headed sections even without MUST (some policies use "will").
    if section.heading and len(section.body) > 80:
        return True
    return False


def _guess_rule_id(text: str) -> str:
    m = RULE_ID_RE.match(text.strip())
    if not m:
        # also accept "R1.1 Returns - ..." mid-line start
        m = re.match(r"^([A-Z]?\d+(?:\.\d+)*)\b", text.strip(), re.IGNORECASE)
    if m:
        return m.group("id" if "id" in m.re.groupindex else 1).upper().replace("SECTION ", "S")
    return ""


def _slug_id(text: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", text.strip())[:40].strip("_").upper()
    return slug or "RULE"


def _split_condition_outcome(body: str) -> tuple[str, str]:
    """Heuristic: sentences with MUST/MUST NOT → outcome; rest → condition."""
    sentences = re.split(r"(?<=[.!?])\s+", body.strip())
    cond, out = [], []
    for s in sentences:
        if re.search(r"\bMUST\b|\bMUST NOT\b|\bshall\b", s, re.IGNORECASE):
            out.append(s)
        else:
            cond.append(s)
    if not out and sentences:
        out = sentences[-1:]
        cond = sentences[:-1]
    return " ".join(cond).strip(), " ".join(out).strip()


def _guess_category(heading: str, body: str) -> str:
    # Prefer the heading — body text often mentions refund/delivery across rules
    # and would otherwise collapse everything into "returns".
    head = heading.lower()
    if any(w in head for w in ("escalat", "manual review", "senior agent")):
        return "global"
    heading_pairs = [
        ("returns", ("return",)),
        ("shipping", ("shipping", "delivery", "carrier", "package", "transit")),
        ("cancellation", ("cancel",)),
        ("warranty", ("warranty", "defect", "misuse")),
        ("billing", ("billing", "charge", "price-match", "price match", "invoice")),
    ]
    for cat, keys in heading_pairs:
        if any(k in head for k in keys):
            return cat

    blob = f"{heading} {body}".lower()
    if any(w in blob for w in ("escalat", "manual review", "senior agent", "flag for")):
        return "global"
    body_pairs = [
        ("cancellation", ("cancel", "cancellation")),
        ("warranty", ("warranty", "defect", "misuse")),
        ("billing", ("billing", "duplicate charge", "price-match", "price match", "invoice")),
        ("shipping", ("shipping", "carrier", "package", "in transit", "promised delivery")),
        ("returns", ("return", "store credit", "final sale")),
    ]
    for cat, keys in body_pairs:
        if any(k in blob for k in keys):
            return cat
    return "global"


def _looks_global(heading: str, body: str) -> bool:
    blob = f"{heading} {body}".lower()
    return any(
        w in blob
        for w in ("escalat", "manual review", "senior agent", "must be flagged", "must not resolve")
    )


def _dedupe_ids(rules: list[PolicyRule]) -> list[PolicyRule]:
    seen: dict[str, int] = {}
    out = []
    for r in rules:
        base = r.id
        n = seen.get(base, 0)
        seen[base] = n + 1
        if n:
            r = r.model_copy(update={"id": f"{base}_{n + 1}"})
        out.append(r)
    return out


# Company-agnostic tokens that indicate a rule conditions on a transaction field.
_DEPENDENCY_TOKENS: dict[str, tuple[str, ...]] = {
    "price": (
        "price", "amount", "total", "value", "cost", "$", "usd", "eur", "gbp",
        "dollar", "refund amount", "order total", "purchase amount",
    ),
    "order_date": (
        "order date", "purchase date", "within", "days of", "day of",
        "return window", "from delivery", "from purchase", "after order",
        "since order", "ordered on", "date of purchase",
    ),
    "status": (
        "status", "delivered", "shipped", "in transit", "cancelled",
        "canceled", "processing", "fulfilled", "fulfilment", "fulfillment",
        "lost", "damaged", "returned", "not shipped",
    ),
    "customer_id": (
        "customer id", "account holder", "registered email", "same customer",
        "customer account", "buyer identity", "verified customer",
    ),
}

# Phrases that suggest the rule is conditional but we couldn't map the field.
_CONDITIONAL_MARKERS = (
    "if ", "when ", "unless ", "provided that", "subject to", "only if",
    "where ", "for orders", "for customers",
)


def _infer_depends_on(condition: str, outcome: str, full: str) -> tuple[list[str], str]:
    """Deterministic structural/token inference of transaction-field dependencies.

    Returns (depends_on, dependency_status). Empty depends_on with status
    "resolved" means the rule is demonstrably unconditional w.r.t. these fields.
    Status "unknown" means the rule looks conditional but no field matched.
    """
    blob = f"{condition}\n{outcome}\n{full}".lower()
    deps: list[str] = []
    for field, tokens in _DEPENDENCY_TOKENS.items():
        if any(tok in blob for tok in tokens):
            deps.append(field)
    if deps:
        # Preserve stable order matching RECOMMENDED_TXN_FIELDS
        order = ["customer_id", "order_date", "price", "status"]
        return [f for f in order if f in deps], "resolved"
    # No field tokens. Is the rule clearly unconditional?
    cond = (condition or "").strip().lower()
    if not cond or not any(m in f"{cond} {full.lower()}" for m in _CONDITIONAL_MARKERS):
        return [], "resolved"
    # Looks conditional but we couldn't map fields.
    return [], "unknown"


def _demo() -> None:
    """Offline self-check — no LLM, uses synthetic markdown."""
    md = """# R1.1 Returns - standard window
A return within 30 days MUST be granted a full refund. Worn items do NOT qualify.

# R6 Escalation - high value
Any case exceeding $200 MUST be escalated to a human senior agent.
"""
    secs = _extract_markdown(md)
    assert len(secs) == 2, secs
    rules = normalize_sections(secs, use_llm_for_changed=False)
    assert rules[0].id.upper().startswith("R1"), rules[0].id
    assert rules[1].category == "global", rules[1]
    assert rules[0].section_hash != rules[1].section_hash
    # hash stability
    assert section_hash(secs[0].text) == secs[0].content_hash
    print("policy_ingest self-check OK")


if __name__ == "__main__":
    _demo()
